import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image
import tkinter as tk
from tkinter import filedialog, messagebox
from docx.enum.section import WD_ORIENT, WD_SECTION

# --- CONFIGURAÇÕES GLOBAIS E PORTABILIDADE ---

try:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
except NameError:
    BASE_DIR = os.getcwd()

LOGO_PATH = os.path.join(BASE_DIR, "logo_porto_2.png")

HEADER_BLUE = "009CDE"
HEADER_TEXT_COLOR = RGBColor(255, 255, 255) 
BODY_TEXT_COLOR = RGBColor(0, 0, 0) 

# ---------------- utilidades OXML ----------------
def set_cell_background(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

# ---------------- formatação global ----------------
def aplicar_formato_run(run, tamanho_pt=None, cor=BODY_TEXT_COLOR):
    run.font.name = "Calibri"
    run.bold = True
    run.font.color.rgb = cor
    if tamanho_pt:
        run.font.size = Pt(tamanho_pt)

def aplicar_formato_paragrafo(p, tamanho_pt=None, cor=BODY_TEXT_COLOR):
    if not p.runs:
        r = p.add_run()
        aplicar_formato_run(r, tamanho_pt, cor)
    else:
        for run in p.runs:
            aplicar_formato_run(run, tamanho_pt, cor)

def padronizar_documento(documento):
    for par in documento.paragraphs:
        for run in par.runs:
            if run.font.color.rgb != HEADER_TEXT_COLOR:
                 aplicar_formato_run(run)
    for tbl in documento.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    for run in par.runs:
                        if run.font.color.rgb != HEADER_TEXT_COLOR:
                             aplicar_formato_run(run)

# ---------------- helpers de imagem ----------------
def file_timestamp(path):
    try:
        ts = os.path.getmtime(path)
        return datetime.fromtimestamp(ts).strftime("%Y-%m-%d %H:%M:%S")
    except:
        return ""

def inserir_imagem_com_titulo(documento, caminho_imagem, largura_inch=6.0):
    nome = os.path.basename(caminho_imagem)

    try:
        documento.add_picture(caminho_imagem, width=Inches(largura_inch))
    except Exception as e:
        p_err = documento.add_paragraph(f"[ERRO AO INSERIR IMAGEM: {nome}] - {e}")
        aplicar_formato_paragrafo(p_err, tamanho_pt=9, cor=RGBColor(255, 0, 0))

# ---------------- função principal ----------------
def criar_documento_porto(caminhos_imagens, cenario_texto, caminho_salvar):
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    tabela = doc.add_table(rows=1, cols=2)
    tabela.autofit = True

    cel_titulo = tabela.rows[0].cells[0]
    cel_logo = tabela.rows[0].cells[1]

    set_cell_background(cel_titulo, HEADER_BLUE)
    set_cell_background(cel_logo, HEADER_BLUE)

    p_t = cel_titulo.paragraphs[0]
    p_t.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run_t = p_t.add_run("EVIDÊNCIAS DE TESTES")
    aplicar_formato_run(run_t, tamanho_pt=20, cor=HEADER_TEXT_COLOR)
    p_t.space_after = Pt(6)

    if os.path.exists(LOGO_PATH):
        p_logo = cel_logo.paragraphs[0]
        p_logo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        try:
            p_logo.add_run().add_picture(LOGO_PATH, width=Inches(1.5))
        except:
            p_logo.add_run("LOGO").bold = True
    else:
        p_logo = cel_logo.paragraphs[0]
        p_logo.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        run_missing = p_logo.add_run("LOGO NÃO ENCONTRADA")
        aplicar_formato_run(run_missing, tamanho_pt=10, cor=HEADER_TEXT_COLOR)

    doc.add_paragraph("")

    h1 = doc.add_paragraph()
    aplicar_formato_run(h1.add_run("Cenário de Teste / Descrição:"), tamanho_pt=14)

    doc.add_paragraph("")

    p_cen = doc.add_paragraph()
    aplicar_formato_run(p_cen.add_run(cenario_texto), tamanho_pt=11)

    doc.add_paragraph("")

    h2 = doc.add_paragraph()
    aplicar_formato_run(h2.add_run("Evidências:"), tamanho_pt=14)

    doc.add_paragraph("")

    if not caminhos_imagens:
        p_empty = doc.add_paragraph()
        aplicar_formato_run(p_empty.add_run("Nenhuma imagem anexada."), tamanho_pt=11)
    else:
        for i, caminho in enumerate(caminhos_imagens):
            p_ev = doc.add_paragraph()
            aplicar_formato_run(p_ev.add_run(f"Evidência {i+1}"), tamanho_pt=12)
            doc.add_paragraph("")
            inserir_imagem_com_titulo(doc, caminho, largura_inch=6.0)
            doc.add_paragraph("")

    doc.save(caminho_salvar)
    return caminho_salvar

# ---------------- GUI ----------------
class AppGUI:
    def __init__(self, master):
        self.master = master
        master.title("Robô Evidências - Modelo Porto Seguro")
        master.minsize(460, 380)

        tk.Label(master, text="Cenário de Teste (cole o texto abaixo):", font=("Calibri", 11, "bold")).pack(pady=(8,0))
        self.texto = tk.Text(master, height=6, width=70, font=("Calibri", 11))
        self.texto.pack(padx=10, pady=6)

        tk.Label(master, text="Imagens selecionadas:", font=("Calibri", 11, "bold")).pack(pady=(6,0))
        self.var_display = tk.StringVar(value="Nenhuma imagem selecionada")
        tk.Label(master, textvariable=self.var_display, bg="lightgray", width=60, anchor="w", wraplength=420).pack(padx=10, pady=6)

        btn_frame = tk.Frame(master)
        btn_frame.pack(pady=6)

        tk.Button(btn_frame, text="Selecionar Imagens", command=self.selecionar_imagens, font=("Calibri", 11, "bold")).grid(row=0, column=0, padx=6)
        tk.Button(btn_frame, text="Gerar Documento (DOCX)", command=self.gerar, bg="green", fg="white", font=("Calibri", 11, "bold")).grid(row=0, column=1, padx=6)

        self.status = tk.Label(master, text="", font=("Calibri", 10, "bold"))
        self.status.pack(pady=(8,6))

        self.caminhos = []

    def selecionar_imagens(self):
        arquivos = filedialog.askopenfilenames(
            title="Selecione imagens",
            filetypes=[("Imagens", "*.png *.jpg *.jpeg"), ("Todos os arquivos", "*.*")],
            initialdir=BASE_DIR
        )

        if arquivos:
            self.caminhos = list(arquivos)
            self.var_display.set(f"{len(self.caminhos)} imagem(ns) selecionada(s).")
        else:
            self.caminhos = []
            self.var_display.set("Nenhuma imagem selecionada")

    def gerar(self):
        texto = self.texto.get("1.0", tk.END).strip()
        if not texto:
            messagebox.showerror("Erro", "O campo do cenário não pode ficar vazio.")
            return

        if not self.caminhos:
            messagebox.showerror("Erro", "Selecione ao menos uma imagem.")
            return

        # ------------------ NOVO: Escolher onde salvar ------------------
        caminho_salvar = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Documento Word", "*.docx")],
            initialfile=f"Relatorio_Evidencias_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )

        if not caminho_salvar:
            return  # usuário cancelou

        self.status.config(text="Gerando documento...", fg="orange")
        self.master.update()

        try:
            caminho_final = criar_documento_porto(self.caminhos, texto, caminho_salvar)
            self.status.config(text=f"Documento criado:\n{caminho_final}", fg="green")
            messagebox.showinfo("Sucesso", f"Documento criado:\n{caminho_final}")
        except Exception as e:
            self.status.config(text=f"Erro: {e}", fg="red")
            messagebox.showerror("Erro", str(e))

if __name__ == "__main__":
    root = tk.Tk()
    app = AppGUI(root)
    root.mainloop()
